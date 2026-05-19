"""Document submission anonymizer.

Takes a set of student document submissions — uploaded individually or inside
a ZIP — and produces an anonymized bundle: every file is renamed to a neutral
code (STUDENT-001.pdf ...) and, where the format allows, personal information
is removed from inside the document. A private name-key CSV maps each code
back to its original filename so grades can be re-attached afterward.

De-identification depth by format:
  - .txt / .md  — structured PII + roster names scrubbed from the text
  - .docx       — structured PII + roster names scrubbed, layout preserved
  - .pdf        — text extracted, redacted, and rebuilt into a clean PDF
                  (visual layout is not preserved — see _redact_pdf)
  - other       — renamed only; contents are not inspected (flagged in summary)

Student names can only be removed reliably when a class roster is supplied,
since names are not otherwise machine-detectable. Without a roster the bundle
still renames files and scrubs structured PII, and the summary says so.

All libraries used here are permissively licensed (pypdf, reportlab: BSD;
python-docx: MIT) so the toolkit carries no copyleft obligations.
"""
from __future__ import annotations

import io
import os
import re
import zipfile
from xml.sax.saxutils import escape as xml_escape

import docx
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

REDACTION_TOKEN = '[REDACTED]'

# Structured PII that regex catches reliably regardless of a roster.
STRUCTURED_PII_PATTERNS = [
    re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'),   # email
    re.compile(r'\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}'),               # phone
    re.compile(r'\b\d{3}-\d{2}-\d{4}\b'),                             # SSN
]

TEXT_EXTENSIONS = {'.txt', '.md'}
MAX_PDF_PAGES = 100

# Junk that macOS / archive tools add; never treated as a submission.
_ARCHIVE_JUNK = ('__MACOSX',)


def build_redaction_terms(roster_names: list[str]) -> list[str]:
    """Expand a roster into literal strings to redact.

    Each full name is kept, plus each name token of 3+ characters (so a
    document that uses only the first or last name is still caught). Sorted
    longest-first so multi-word names redact before their individual tokens.
    """
    terms: set[str] = set()
    for raw in roster_names:
        name = raw.strip()
        if len(name) < 2:
            continue
        terms.add(name)
        for token in name.split():
            if len(token) >= 3:
                terms.add(token)
    return sorted(terms, key=len, reverse=True)


def _make_term_patterns(terms: list[str]) -> list[re.Pattern]:
    """Compile case-insensitive, word-boundaried patterns for literal terms."""
    return [
        re.compile(r'(?<!\w)' + re.escape(term) + r'(?!\w)', re.IGNORECASE)
        for term in terms
    ]


def _redact_text(text: str, term_patterns: list[re.Pattern]) -> tuple[str, int]:
    """Replace roster terms and structured PII in a plain string."""
    count = 0
    out = text
    for pattern in term_patterns:
        out, n = pattern.subn(REDACTION_TOKEN, out)
        count += n
    for pattern in STRUCTURED_PII_PATTERNS:
        out, n = pattern.subn(REDACTION_TOKEN, out)
        count += n
    return out, count


def _redact_txt(file_bytes: bytes, term_patterns: list[re.Pattern]) -> tuple[bytes, int]:
    text = file_bytes.decode('utf-8', errors='replace')
    redacted, count = _redact_text(text, term_patterns)
    return redacted.encode('utf-8'), count


def _iter_docx_paragraphs(document: "docx.document.Document"):
    """Yield every paragraph in a docx: body, tables, and section headers/footers."""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _redact_paragraph(paragraph, term_patterns: list[re.Pattern]) -> int:
    """Redact a docx paragraph, joining runs so names split across runs are caught.

    When a paragraph contains PII its runs are collapsed into the first run.
    Formatting is only flattened on paragraphs that actually held PII.
    """
    runs = paragraph.runs
    if not runs:
        return 0
    joined = ''.join(run.text for run in runs)
    if not joined:
        return 0
    redacted, count = _redact_text(joined, term_patterns)
    if count == 0:
        return 0
    runs[0].text = redacted
    for run in runs[1:]:
        run.text = ''
    return count


def _redact_docx(file_bytes: bytes, term_patterns: list[re.Pattern]) -> tuple[bytes, int]:
    document = docx.Document(io.BytesIO(file_bytes))
    count = 0
    for paragraph in _iter_docx_paragraphs(document):
        count += _redact_paragraph(paragraph, term_patterns)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), count


def _text_to_pdf(text: str) -> bytes:
    """Build a simple single-column PDF from plain text."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    style = getSampleStyleSheet()['Normal']
    flowables: list = []
    for line in text.split('\n'):
        stripped = line.strip()
        if stripped:
            flowables.append(Paragraph(xml_escape(stripped), style))
        else:
            flowables.append(Spacer(1, 6))
    doc.build(flowables or [Spacer(1, 6)])
    return buffer.getvalue()


def _redact_pdf(
    file_bytes: bytes, term_patterns: list[re.Pattern],
) -> tuple[bytes, int, str | None, bool]:
    """Redact a PDF by extracting its text, scrubbing it, and rebuilding a clean PDF.

    The original visual layout is not preserved — the anonymized PDF is
    regenerated from the redacted text. This is deliberate: rebuilding from
    scratch guarantees no hidden original content (off-page text, annotations,
    metadata) survives, which is the safest outcome for de-identification.
    A scanned/image-only or locked PDF yields no text; the original is kept
    unchanged and flagged so the faculty member redacts it manually.
    """
    reader = PdfReader(io.BytesIO(file_bytes))
    if reader.is_encrypted:
        try:
            reader.decrypt('')
        except Exception:
            pass

    total_pages = len(reader.pages)
    pages_to_read = min(total_pages, MAX_PDF_PAGES)
    chunks: list[str] = []
    for index in range(pages_to_read):
        try:
            chunks.append(reader.pages[index].extract_text() or '')
        except Exception:
            chunks.append('')

    raw_text = '\n\n'.join(c for c in chunks if c.strip())
    if not raw_text.strip():
        return (
            file_bytes, 0,
            'No readable text — the PDF may be scanned, image-only, or '
            'password-protected. Original kept; redact it manually before sharing.',
            False,
        )

    redacted, count = _redact_text(raw_text, term_patterns)
    note = None
    if total_pages > MAX_PDF_PAGES:
        note = f'Only the first {MAX_PDF_PAGES} of {total_pages} pages were processed.'
    return _text_to_pdf(redacted), count, note, True


def _process_one(
    data: bytes,
    ext: str,
    term_patterns: list[re.Pattern],
) -> tuple[bytes, int, str | None, bool]:
    """Anonymize a single document. Returns (bytes, redactions, note, content_redacted).

    Any processing failure degrades safely to rename-only so one malformed
    file never aborts the whole batch.
    """
    try:
        if ext in TEXT_EXTENSIONS:
            out, count = _redact_txt(data, term_patterns)
            return out, count, None, True
        if ext == '.docx':
            out, count = _redact_docx(data, term_patterns)
            return out, count, None, True
        if ext == '.pdf':
            return _redact_pdf(data, term_patterns)
        return (
            data, 0,
            f'Format "{ext or "unknown"}" not supported for content redaction — renamed only.',
            False,
        )
    except Exception as exc:  # noqa: BLE001 - degrade to rename-only on any error
        return data, 0, f'Could not process contents ({exc}) — file renamed only.', False


def _match_roster(filename: str, roster_names: list[str]) -> str:
    """Best-effort: link a file to a roster name when the name is in the filename."""
    norm_file = re.sub(r'[^a-z0-9]', '', filename.lower())
    matches = [
        name for name in roster_names
        if (norm := re.sub(r'[^a-z0-9]', '', name.lower())) and norm in norm_file
    ]
    return matches[0] if len(matches) == 1 else ''


def read_archive(zip_bytes: bytes) -> list[tuple[str, bytes]]:
    """Extract (filename, bytes) pairs from a ZIP, skipping folders and junk."""
    files: list[tuple[str, bytes]] = []
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as archive:
        for info in archive.infolist():
            if info.is_dir():
                continue
            if any(junk in info.filename for junk in _ARCHIVE_JUNK):
                continue
            base = info.filename.rsplit('/', 1)[-1]
            if not base or base.startswith('.'):
                continue
            files.append((base, archive.read(info)))
    return files


def anonymize_submissions(
    files: list[tuple[str, bytes]], roster_names: list[str],
) -> dict:
    """Anonymize a batch of submissions into a ZIP plus a private name key.

    Returns ``zip_bytes`` (the anonymized bundle for the AI tool),
    ``name_key_csv`` (private mapping of code to original file), and a
    ``summary`` describing what was redacted per file.
    """
    if not files:
        raise ValueError('No documents were provided to anonymize.')

    clean_roster = [n.strip() for n in roster_names if n.strip()]
    terms = build_redaction_terms(clean_roster)
    term_patterns = _make_term_patterns(terms)

    code_width = max(3, len(str(len(files))))
    per_file: list[dict] = []
    name_key_rows: list[tuple[str, str, str]] = []

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as bundle:
        for index, (orig_name, data) in enumerate(files, start=1):
            code = f'STUDENT-{index:0{code_width}d}'
            ext = os.path.splitext(orig_name)[1].lower()
            out_bytes, redactions, note, content_redacted = _process_one(
                data, ext, term_patterns,
            )
            bundle.writestr(f'{code}{ext}', out_bytes)
            name_key_rows.append((code, orig_name, _match_roster(orig_name, clean_roster)))
            per_file.append({
                'code': code,
                'originalName': orig_name,
                'format': ext or 'unknown',
                'contentRedacted': content_redacted,
                'redactionCount': redactions,
                'note': note,
            })

    key_buffer = io.StringIO()
    key_buffer.write('submission_id,Original Filename,Matched Student\r\n')
    for code, orig_name, matched in name_key_rows:
        key_buffer.write(_csv_row(code, orig_name, matched))

    total_files = len(files)
    return {
        'zip_bytes': zip_buffer.getvalue(),
        'name_key_csv': key_buffer.getvalue(),
        'summary': {
            'fileCount': total_files,
            'codeRange': (
                f'STUDENT-{1:0{code_width}d} through STUDENT-{total_files:0{code_width}d}'
            ),
            'rosterNamesProvided': len(clean_roster),
            'totalRedactions': sum(f['redactionCount'] for f in per_file),
            'contentRedactedCount': sum(1 for f in per_file if f['contentRedacted']),
            'renamedOnlyCount': sum(1 for f in per_file if not f['contentRedacted']),
            'perFile': per_file,
        },
    }


def _csv_row(*values: str) -> str:
    """Minimal RFC-4180 CSV row writer (quote when needed)."""
    cells = []
    for value in values:
        if any(c in value for c in (',', '"', '\n', '\r')):
            cells.append('"' + value.replace('"', '""') + '"')
        else:
            cells.append(value)
    return ','.join(cells) + '\r\n'

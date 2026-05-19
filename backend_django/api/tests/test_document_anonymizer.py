"""Tests for the document submission anonymizer: service + endpoint."""
import base64
import io
import zipfile

import docx
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.test import TestCase, Client
from django.contrib.auth.models import User
from django.core.files.uploadedfile import SimpleUploadedFile

from api.models import UserProfile
from api.services.document_anonymizer import (
    anonymize_submissions,
    build_redaction_terms,
    read_archive,
)


def _make_pdf(text: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.drawString(72, 700, text)
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return '\n'.join(page.extract_text() or '' for page in reader.pages)


def _make_docx(paragraphs: list[str]) -> bytes:
    document = docx.Document()
    for para in paragraphs:
        document.add_paragraph(para)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_text(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    return '\n'.join(p.text for p in document.paragraphs)


def _make_zip(members: dict[str, bytes]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w') as archive:
        for name, data in members.items():
            archive.writestr(name, data)
    return buffer.getvalue()


class RedactionTermsTest(TestCase):
    def test_expands_full_name_and_tokens(self) -> None:
        terms = build_redaction_terms(['Jane Doe', 'Bo Ng'])
        self.assertIn('Jane Doe', terms)
        self.assertIn('Jane', terms)
        self.assertIn('Doe', terms)
        # The full roster entry is always kept, even when short...
        self.assertIn('Bo Ng', terms)
        # ...but short individual tokens (< 3 chars) are not redacted alone.
        self.assertNotIn('Bo', terms)
        self.assertNotIn('Ng', terms)
        # Longest-first ordering so multi-word names redact before tokens.
        self.assertEqual(terms[0], 'Jane Doe')


class AnonymizeSubmissionsServiceTest(TestCase):
    def test_txt_redacts_roster_name_and_structured_pii(self) -> None:
        content = b'Essay by Jane Doe. Contact: jane@usf.edu or 813-555-1234.'
        result = anonymize_submissions([('jane_doe_essay.txt', content)], ['Jane Doe'])
        with zipfile.ZipFile(io.BytesIO(result['zip_bytes'])) as bundle:
            names = bundle.namelist()
            self.assertEqual(names, ['STUDENT-001.txt'])
            out = bundle.read('STUDENT-001.txt').decode('utf-8')
        self.assertNotIn('Jane Doe', out)
        self.assertNotIn('jane@usf.edu', out)
        self.assertNotIn('813-555-1234', out)
        self.assertIn('[REDACTED]', out)

    def test_docx_redacts_name_inside_document(self) -> None:
        data = _make_docx(['Submitted by Jane Doe', 'My argument is that...'])
        result = anonymize_submissions([('essay.docx', data)], ['Jane Doe'])
        with zipfile.ZipFile(io.BytesIO(result['zip_bytes'])) as bundle:
            out = bundle.read('STUDENT-001.docx')
        text = _docx_text(out)
        self.assertNotIn('Jane Doe', text)
        self.assertIn('[REDACTED]', text)
        self.assertIn('My argument is that', text)

    def test_pdf_redacts_name_in_place(self) -> None:
        data = _make_pdf('Final Essay by Jane Doe for MIS 4123')
        result = anonymize_submissions([('essay.pdf', data)], ['Jane Doe'])
        with zipfile.ZipFile(io.BytesIO(result['zip_bytes'])) as bundle:
            out = bundle.read('STUDENT-001.pdf')
        self.assertNotIn('Jane Doe', _pdf_text(out))

    def test_unsupported_format_is_renamed_only(self) -> None:
        result = anonymize_submissions([('photo.png', b'\x89PNG fake')], [])
        entry = result['summary']['perFile'][0]
        self.assertFalse(entry['contentRedacted'])
        self.assertEqual(result['summary']['renamedOnlyCount'], 1)
        with zipfile.ZipFile(io.BytesIO(result['zip_bytes'])) as bundle:
            self.assertEqual(bundle.namelist(), ['STUDENT-001.png'])

    def test_name_key_maps_codes_to_filenames(self) -> None:
        files = [('jane_doe.txt', b'x'), ('bob_lee.txt', b'y')]
        result = anonymize_submissions(files, ['Jane Doe', 'Bob Lee'])
        key = result['name_key_csv']
        self.assertIn('submission_id,Original Filename,Matched Student', key)
        self.assertIn('STUDENT-001,jane_doe.txt,Jane Doe', key)
        self.assertIn('STUDENT-002,bob_lee.txt,Bob Lee', key)

    def test_empty_batch_rejected(self) -> None:
        with self.assertRaises(ValueError):
            anonymize_submissions([], ['Jane Doe'])

    def test_read_archive_skips_folders_and_junk(self) -> None:
        zip_bytes = _make_zip({
            'essays/': b'',
            'essays/jane.txt': b'content',
            '__MACOSX/._jane.txt': b'junk',
            '.DS_Store': b'junk',
        })
        files = read_archive(zip_bytes)
        self.assertEqual(len(files), 1)
        self.assertEqual(files[0][0], 'jane.txt')


class AnonymizeDocumentsEndpointTest(TestCase):
    def setUp(self) -> None:
        self.client = Client()
        self.user = User.objects.create_user(
            username='doc@usf.edu', email='doc@usf.edu',
            password='testpass123', first_name='Doc',
        )
        UserProfile.objects.create(user=self.user, role='pi')
        self.url = '/api/verify/anonymize-documents'

    def _login(self) -> None:
        self.client.login(username='doc@usf.edu', password='testpass123')

    def test_unauth_blocked(self) -> None:
        response = self.client.post(self.url, {})
        self.assertEqual(response.status_code, 401)

    def test_no_files_rejected(self) -> None:
        self._login()
        response = self.client.post(self.url, {'roster': 'Jane Doe'})
        self.assertEqual(response.status_code, 400)

    def test_non_zip_archive_rejected(self) -> None:
        self._login()
        bad = SimpleUploadedFile('subs.txt', b'not a zip', content_type='text/plain')
        response = self.client.post(self.url, {'archive': bad})
        self.assertEqual(response.status_code, 400)

    def test_individual_files_anonymized(self) -> None:
        self._login()
        f1 = SimpleUploadedFile('jane.txt', b'Essay by Jane Doe', content_type='text/plain')
        f2 = SimpleUploadedFile('bob.txt', b'Essay by Bob Lee', content_type='text/plain')
        response = self.client.post(self.url, {'files': [f1, f2], 'roster': 'Jane Doe\nBob Lee'})
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body['summary']['fileCount'], 2)
        zip_bytes = base64.b64decode(body['anonymizedZipBase64'])
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as bundle:
            self.assertEqual(sorted(bundle.namelist()), ['STUDENT-001.txt', 'STUDENT-002.txt'])
            self.assertNotIn(b'Jane Doe', bundle.read('STUDENT-001.txt'))

    def test_zip_archive_anonymized(self) -> None:
        self._login()
        zip_bytes = _make_zip({
            'jane_essay.txt': b'By Jane Doe, jane@usf.edu',
            'bob_essay.txt': b'By Bob Lee',
        })
        archive = SimpleUploadedFile('submissions.zip', zip_bytes, content_type='application/zip')
        response = self.client.post(self.url, {'archive': archive, 'roster': 'Jane Doe, Bob Lee'})
        self.assertEqual(response.status_code, 200)
        body = response.json()
        self.assertEqual(body['summary']['fileCount'], 2)
        self.assertIn('submission_id,Original Filename', body['nameKeyCsv'])
